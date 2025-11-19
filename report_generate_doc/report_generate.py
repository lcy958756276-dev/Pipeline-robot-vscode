#!/usr/bin/env python
"""
作用：
    接收管道检测相关的 ROS 消息，将检测到的缺陷信息（含图像）写入 Word 文档用于报告生成。

参数：
    无（通过 ROS 订阅接收数据）。

返回：
    无（运行在 ROS 节点内，执行 I/O 操作保存文档）。
"""

import rospy
from pipeline_robot.msg import detection_new
from nav_msgs.msg import Odometry
from cv_bridge import CvBridge, CvBridgeError
import cv2
from docx import Document
from docx.shared import Inches
import os
import datetime
import xml.etree.ElementTree as ET
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT

class ImageSaver:
    """
    作用：
        订阅缺陷检测与里程计话题，接收缺陷信息并把带图像的缺陷条目写入 Word 文档（报告）。

    参数：
        无（类在实例化时会初始化 ROS 订阅者和内部状态）。

    返回：
        无（实例化后在 ROS spin 循环中运行）。
    """

    def __init__(self):
        """
        作用：
            初始化 ImageSaver 对象，设置 CvBridge、Document、输出路径、订阅话题等初始状态。

        参数：
            无

        返回：
            None
        """
        rospy.loginfo("===================================enter report generate========================")
        self.bridge = CvBridge()
        self.document = Document()
        current_time = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        filename = f"detection55_{current_time}.docx"
        self.doc_path = f'/home/ralab/pipeline-inspection-robot-topnx/src/pipeline_robot/report_test/{filename}'
        # self.doc_path = '/home/jetson/pipeline-inspection-robot/src/pipeline_robot/report/3.docx'
        self.detected_defects = []  # 存储已检测到的缺陷
        self.current_odom_x = 0.0  # 初始化当前odom x值
        self.is_stop=True

        # 订阅detection_info和odom话题
        rospy.Subscriber("/_detection_info", detection_new, self.detection_callback)
        # rospy.Subscriber("/odometry/filtered", Odometry, self.odometry_callback)
        rospy.Subscriber("/cable/odom", Odometry, self.odom_callback)
        rospy.Subscriber("/odom", Odometry, self.state_callback)

    def detection_callback(self, data):
        """
        作用：
            接收来自 /_detection_info 的检测消息，记录缺陷并触发保存函数。

        参数：
            data (detection_new): 来自检测模块的消息，包含 class_name、point_stamped、image_data、size 等字段。

        返回：
            None
        """
        rospy.loginfo("recv detection data: %s",data)
        # if not self.is_defect_unuseful(data):
        self.detected_defects.append((data.class_name, data.point_stamped.point.x+self.current_odom_x,  data.point_stamped.point.y,  data.point_stamped.point.z))
        self.save_defect_info(data)


    # def odometry_callback(self, data):
        # 更新当前odom x值
        # self.current_odom_x = data.pose.pose.position.x
        # if abs(data.twist.twist.linear.x)<0.01:
        #     self.is_stop = True
        # else:
        #     self.is_stop = False
    def odom_callback(self, data):
        """
        作用：
            接收里程计数据并更新内部的当前全局 x 位置（用于将检测点转换为全局坐标）。

        参数：
            data (Odometry): 里程计消息，包含位姿信息。

        返回：
            None
        """
        # 更新当前odom x值
        self.current_odom_x = data.pose.pose.position.x
        
    def state_callback(self, data):
        """
        作用：
            根据里程计的线速度判断机器人是否处于静止状态，并更新内部标志 is_stop。

        参数：
            data (Odometry): 里程计消息，包含速度信息。

        返回：
            None
        """
        
        if abs(data.twist.twist.linear.x)<0.01:
            self.is_stop = True
        else:
            self.is_stop = False

    def is_defect_unuseful(self, data, threshold=0.5):
        """
        作用：
            判断当前检测到的缺陷是否应被忽略（例如异常点、机器人未静止、距离过远或与已记录缺陷重复）。

        参数：
            data (detection_new): 检测消息，含坐标与类别等信息。
            threshold (float): 判定与已记录缺陷重复的距离阈值，默认 0.5。

        返回：
            bool: 若返回 True 表示该缺陷不需要记录，False 表示应记录。
        """
    
    # if self.is_stop!=0:
    #     return True
        if abs(data.point_stamped.point.x-0)<0.01:
            rospy.loginfo("当前目标点异常，缺陷类型: {},X: {}, Y: {}, Z: {}".format(data.class_name,data.point_stamped.point.x, data.point_stamped.point.y, data.point_stamped.point.z))
            return True
        if self.is_stop!=True:
            rospy.loginfo("机器人未静止")
            return True
        if data.point_stamped.point.x>1.4:
            rospy.loginfo("检测到缺陷类型: {}距离过远为:{}，暂不记录".format(data.class_name,data.point_stamped.point.x))
            return True
        for defect in self.detected_defects:
            if data.class_name == defect[0]:
                # 将odom的x值考虑在内
                dx, dy, dz = abs((data.point_stamped.point.x + self.current_odom_x) - defect[1]), abs(data.point_stamped.point.y - defect[2]), abs(data.point_stamped.point.z - defect[3])
                if dx < threshold :
                    return True
        return False

    def set_table_font_size(self,table, font_size_pt):
        """
        作用：
            将指定表格内所有单元格文本的字体大小设置为指定磅数。

        参数：
            table: docx 的 table 对象。
            font_size_pt (int/float): 字体大小（磅）。

        返回：
            None
        """
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(font_size_pt)

    def save_defect_info(self, data):
        """
        作用：
            将单条缺陷信息写入 Word 文档：包含将 ROS 图像转换成文件、在文档中创建表格并插入图片和缺陷信息，最后保存文档。

        参数：
            data (detection_new): 来自检测模块的消息，包含图像、类别、位置信息与尺寸等字段。

        返回：
            None
        """
        try:
            cv_image = self.bridge.imgmsg_to_cv2(data.image_data, "bgr8")
        except CvBridgeError as e:
            rospy.logerr("CvBridge Error: {0}".format(e))
            return

        temp_image_path = '/home/ralab/pipeline-inspection-robot-topnx/src/pipeline_robot/report/image.jpg'
        cv2.imwrite(temp_image_path, cv_image)
        self.document.add_paragraph(' ')
        table = self.document.add_table(rows=11, cols=5)
        # 尝试将表格对齐到文档的最左端
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # 设置表格的左缩进为0，尝试使表格对齐到最左端
        table_left_indent = OxmlElement('w:tblInd')
        table_left_indent.set(qn('w:w'), "0")  # 缩进值，单位为dxa（twentieths of a point）
        table_left_indent.set(qn('w:type'), 'dxa')
        table._tblPr.append(table_left_indent)
 


        for row in range(11):
            for col in range(5):
                table.cell(row, col).text = u''
        # column_widths = [Inches(0.5), Inches(0.5), Inches(0.5), Inches(0.5),Inches(1)]
        # for i, column in enumerate(table.columns):
        #     if i < len(column_widths):  # 仅调整前5列
        #         self.set_column_width(column, column_widths[i])
        # length = 2
        # for col in range(0, 5):
        #     for row in range(0, 11):
        #         table.cell(row, col).width = Cm(length)
        # table.cell(0,0).width=Cm(2)
        # table.cell(1,0).width=Cm(2)
        # table.cell(2,0).width=Cm(2)
        # table.cell(3,0).width=Cm(2)
        # table.cell(4,0).width=Cm(2)
        # table.cell(5,0).width=Cm(2)
        # table.cell(6,0).width=Cm(2)
        # table.cell(7,0).width=Cm(2)
        # table.cell(8,0).width=Cm(2)
        # table.cell(9,0).width=Cm(2)
        # table.cell(10,0).width=Cm(2)
        col1 = table.columns[0]
        col1.width = Inches(0.7)
        col2 = table.columns[1]
        col2.width = Inches(0.7)
        col3 = table.columns[2]
        col3.width = Inches(0.7)
        col4 = table.columns[3]
        col4.width = Inches(0.7)
        col5 = table.columns[4]
        col5.width = Inches(4)


        for row in range(7):
            table.cell(row, 1).merge(table.cell(row, 3))


        table.cell(10, 1).merge(table.cell(10, 3))


        table.cell(0, 4).merge(table.cell(10, 4))
        
        #existed
        table.cell(0, 0).text = u'管段编号 '
        table.cell(1, 0).text = u'水流方向 '
        table.cell(2, 0).text = u'管道材质 '
        table.cell(3, 0).text = u'管道直径 '
        table.cell(4, 0).text = u'视频文件 '
        table.cell(5, 0).text = u'缺陷名称 '
        table.cell(6, 0).text = u'检测日期 '
        table.cell(7, 0).text = u'缺陷登记 '
        table.cell(8, 0).text = u'时钟表示 '
        table.cell(9, 0).text = u'管端埋深 '
        table.cell(10, 0).text = u'缺陷描述 '
        table.cell(9, 2).text = u'管端长度 '
        table.cell(8, 2).text = u'缺陷长度 '
        table.cell(7, 2).text = u'缺陷距离 '
        table.cell(0, 1).text = u''
        table.cell(0, 4).text = u''
        table.cell(1, 1).text = u''
        table.cell(2, 1).text = u''
        table.cell(3, 1).text = u''
        table.cell(4, 1).text = u''
        table.cell(10, 1).text = u''
        table.cell(8, 3).text = data.size
        table.cell(5, 1).text = u'{} '.format(data.class_name)
        table.cell(6, 1).text = u'{} '.format(datetime.datetime.now().strftime("%Y-%m-%d"))
        table.cell(7,3).text = u'{:.3f} '.format(data.point_stamped.point.x+self.current_odom_x)

        paragraph = table.cell(0,4).paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(temp_image_path, width=Inches(4.0))

        # self.document.add_picture(temp_image_path, width=Inches(4.0))
        # self.document.add_paragraph('缺陷类型: {}'.format(data.class_name))
        
        
        # 在位置信息中加入odom x值
        # self.document.add_paragraph('当前机器人在管道全局位置信息:  X: {}'.format(self.current_odom_x))
        # self.document.add_paragraph('当前缺陷距离机器人的位置信息:   X: {}, Y: {}, Z: {}'.format(data.point_stamped.point.x, data.point_stamped.point.y, data.point_stamped.point.z))
        # self.document.add_paragraph('该缺陷在管道全局位置信息:  X: {}, Y: {}, Z: {}'.format(data.point_stamped.point.x+self.current_odom_x, data.point_stamped.point.y, data.point_stamped.point.z))
        self.set_table_font_size(table, 9)
        self.document.save(self.doc_path)
        # os.remove(temp_image_path)

if __name__ == '__main__':
    rospy.init_node('report_generate', anonymous=True)
    image_saver = ImageSaver()
    rospy.spin()
